import os
import sys
import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn

def export_to_docx(json_data_path: str, output_path: str):
    with open(json_data_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    title = data.get("title", "手写文章数字化结果")
    pages = data.get("pages", [])

    doc = Document()

    # Set standard page margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Set default font for Chinese
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.styles['Normal'].font.size = Pt(12)

    # Document Title
    heading = doc.add_heading(title, level=0)
    heading.style.font.name = 'Times New Roman'
    heading.style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    # Add Pages Content
    for idx, page in enumerate(pages):
        if len(pages) > 1:
            page_title = doc.add_heading(f"第 {idx + 1} 页", level=1)
            page_title.style.font.name = 'Times New Roman'
            page_title.style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        text = page.get("editedText", "")
        if not text:
            text = page.get("rawText", "")

        lines = text.split("\n")
        for line in lines:
            trimmed = line.strip()
            if trimmed:
                p = doc.add_paragraph(trimmed)
                p.paragraph_format.line_spacing = 1.35
                p.paragraph_format.space_after = Pt(6)

        if idx < len(pages) - 1:
            doc.add_page_break()

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    print(f"Successfully exported DOCX to {output_path}")

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python export_docx.py <input_json> <output_docx>")
        sys.exit(1)
    export_to_docx(sys.argv[1], sys.argv[2])
